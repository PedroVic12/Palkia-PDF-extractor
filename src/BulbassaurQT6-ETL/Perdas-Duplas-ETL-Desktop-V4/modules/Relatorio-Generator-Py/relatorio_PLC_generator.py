import sys
import os
import time
import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.shared import Inches, RGBColor
from docx2pdf import convert
from tqdm import tqdm

# --- IMPORTS DO PYSIDE6 ---
from PySide6.QtWidgets import (QApplication, QMainWindow, QWidget, QVBoxLayout, 
                               QHBoxLayout, QLabel, QLineEdit, QPushButton, 
                               QTextEdit, QFileDialog, QMessageBox, QFrame, QSplitter)
from PySide6.QtCore import Qt, QThread, Signal, QObject
from PySide6.QtGui import QFont, QColor, QTextCursor, QTextCharFormat

# ============================================================================
# 1. LÓGICA DE NEGÓCIO (BACKEND)
# ============================================================================

def adicionar_tabela_teste(doc: Document, num_linhas: int = 5, num_colunas: int = 5):
    print(f"\n📊 [INFO] Adicionando tabela de teste ({num_linhas}x{num_colunas})...")
    
    doc.add_heading('Tabela de Teste - Formatação', level=2)
    doc.add_paragraph("Esta é uma tabela de teste para validar a formatação dos relatórios:")
    
    table = doc.add_table(rows=num_linhas + 1, cols=num_colunas)
    table.style = 'Table Grid'
    
    header_cells = table.rows[0].cells
    colunas_nomes = [f"Coluna {i+1}" for i in range(num_colunas)]
    
    for j, col_name in enumerate(colunas_nomes):
        cell = header_cells[j]
        cell.text = col_name
        paragraph = cell.paragraphs[0]
        paragraph.style = 'Heading 3'
        for run in paragraph.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
        
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls
        cell._element.get_or_add_tcPr().append(
            parse_xml(r'<w:shd {} w:fill="4472C4"/>'.format(nsdecls('w')))
        )
    
    for i in range(num_linhas):
        row_cells = table.rows[i + 1].cells
        for j in range(num_colunas):
            row_cells[j].text = f"Dado {i+1}-{j+1}"
    
    print(f"✅ [SUCESSO] Tabela de teste adicionada!")

def processar_relatorio(template_path, excel_path, output_docx, output_pdf, apenas_teste=False):
    """Função mestre que orquestra a criação do relatório."""
    
    # Validações básicas
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template não encontrado: {template_path}")
    
    # Se não for apenas teste, precisa do Excel
    if not apenas_teste and not os.path.exists(excel_path):
        raise FileNotFoundError(f"Excel não encontrado: {excel_path}")

    # Garante caminhos absolutos
    template_path = os.path.abspath(template_path)
    if excel_path: excel_path = os.path.abspath(excel_path)
    output_docx = os.path.abspath(output_docx)
    if output_pdf: output_pdf = os.path.abspath(output_pdf)

    print(f"📂 [IO] Abrindo template: {os.path.basename(template_path)}")
    doc = Document(template_path)
    
    if apenas_teste:
        print("🧪 [MODE] Modo de Validação: Apenas gerando tabela de teste...")
        adicionar_tabela_teste(doc)
        doc.save(output_docx)
        print(f"💾 [SAVE] Documento de teste salvo em: {os.path.basename(output_docx)}")
        return

    # --- MODO COMPLETO ---
    
    # 1. Adiciona tabela de teste (Opcional, mantido conforme seu código)
    adicionar_tabela_teste(doc)
    
    # 2. Nova seção para dados
    nova_secao = doc.add_section(WD_SECTION.NEW_PAGE)
    nova_secao.top_margin = Inches(1)
    nova_secao.bottom_margin = Inches(1)
    doc.add_heading('Relatório de Perdas Duplas (Dados do Excel)', level=1)
    doc.add_paragraph("Abaixo estão os dados extraídos da planilha Excel:")

    # 3. Carrega Excel
    print(f"📊 [DATA] Lendo Excel: {os.path.basename(excel_path)}")
    df = pd.read_excel(excel_path)
    rows, cols = df.shape
    
    table = doc.add_table(rows + 1, cols)
    table.style = 'Table Grid'

    # Cabeçalho
    for j, column_name in enumerate(df.columns):
        table.cell(0, j).text = str(column_name)

    # Dados com barra de progresso customizada para GUI
    print(f"📝 [WORD] Preenchendo tabela com {rows} linhas...")
    
    # Nota: O tqdm padrão imprime caracteres de controle que podem sujar o log da GUI
    # Vamos usar um loop simples com print de porcentagem para ficar mais limpo na GUI
    for i in range(rows):
        # Feedback visual a cada 10%
        if i % (max(1, rows // 10)) == 0:
            perc = int((i / rows) * 100)
            print(f"   ... Processando linha {i}/{rows} ({perc}%)")
            
        for j in range(cols):
            val = df.values[i, j]
            table.cell(i + 1, j).text = str(val) if pd.notna(val) else ""

    doc.save(output_docx)
    print(f"💾 [SAVE] Word salvo: {os.path.basename(output_docx)}")

    # 4. Conversão PDF
    print(f"🔄 [PDF] Iniciando conversão (Aguarde o Word)...")
    start_time = time.time()
    convert(output_docx, output_pdf)
    end_time = time.time()
    print(f"✅ [DONE] Conversão concluída em {end_time - start_time:.2f}s")
    print(f"📄 [FILE] Arquivo PDF: {os.path.basename(output_pdf)}")

# ============================================================================
# 2. INFRAESTRUTURA VISUAL (FRONTEND)
# ============================================================================

class EmittingStream(QObject):
    """Redireciona print() para o Widget QTextEdit"""
    textWritten = Signal(str)
    def write(self, text): self.textWritten.emit(str(text))
    def flush(self): pass

class WorkerThread(QThread):
    """Thread para rodar a lógica sem travar a tela"""
    finished_signal = Signal()
    error_signal = Signal(str)

    def __init__(self, t_path, e_path, o_word, o_pdf, mode):
        super().__init__()
        self.t_path = t_path
        self.e_path = e_path
        self.o_word = o_word
        self.o_pdf = o_pdf
        self.mode = mode # 1 = Completo, 2 = Validação

    def run(self):
        try:
            apenas_teste = (self.mode == 2)
            processar_relatorio(self.t_path, self.e_path, self.o_word, self.o_pdf, apenas_teste)
            self.finished_signal.emit()
        except Exception as e:
            self.error_signal.emit(str(e))

class RichConsoleWidget(QTextEdit):
    """Widget que simula um terminal moderno"""
    def __init__(self):
        super().__init__()
        self.setReadOnly(True)
        self.setStyleSheet("""
            QTextEdit {
                background-color: #1e1e1e;
                color: #d4d4d4;
                font-family: 'Consolas', 'Courier New', monospace;
                font-size: 14px;
                border: 1px solid #333;
                border-radius: 5px;
                padding: 10px;
            }
        """)
    
    def append_colored(self, text):
        # Simples parser para colorir logs baseados em keywords
        cursor = self.textCursor()
        cursor.movePosition(QTextCursor.End)
        
        format = QTextCharFormat()
        
        if "❌" in text or "[ERRO]" in text:
            format.setForeground(QColor("#ff5555")) # Vermelho
        elif "✅" in text or "[SUCESSO]" in text or "[DONE]" in text:
            format.setForeground(QColor("#50fa7b")) # Verde
        elif "📊" in text or "[DATA]" in text:
            format.setForeground(QColor("#f1fa8c")) # Amarelo
        elif "🔄" in text or "[PDF]" in text:
            format.setForeground(QColor("#bd93f9")) # Roxo
        else:
            format.setForeground(QColor("#d4d4d4")) # Cinza claro padrão
            
        cursor.setCharFormat(format)
        cursor.insertText(text)
        self.setTextCursor(cursor)
        self.ensureCursorVisible()

class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        
        # Redirecionamento de Print
        self.sys_stdout = sys.stdout
        self.emitter = EmittingStream()
        self.emitter.textWritten.connect(self.on_print)
        sys.stdout = self.emitter

        self.setup_ui()

    def setup_ui(self):
        self.setWindowTitle("Sistema de Relatórios ONS - Perdas Duplas")
        self.resize(1200, 700)
        self.setStyleSheet("background-color: #f0f0f0;")

        # Widget Central
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        main_layout = QHBoxLayout(central_widget)

        # --- LADO ESQUERDO: CONTROLES (GUI) ---
        left_panel = QFrame()
        left_panel.setFixedWidth(500)
        left_panel.setStyleSheet("""
            QFrame { background-color: white; border-radius: 10px; padding: 10px; }
            QLabel { font-size: 16px; font-weight: bold; color: #333; font-family: 'Segoe UI'; }
            QLineEdit { font-size: 14px; padding: 8px; border: 1px solid #ccc; border-radius: 4px; }
            QPushButton { font-size: 14px; padding: 10px; border-radius: 5px; font-weight: bold; }
        """)
        left_layout = QVBoxLayout(left_panel)
        left_layout.setSpacing(15)

        # Título
        lbl_header = QLabel("⚙️ Configurações do Processo")
        lbl_header.setStyleSheet("font-size: 22px; color: #0088cc; border-bottom: 2px solid #eee; padding-bottom: 10px;")
        left_layout.addWidget(lbl_header)

        # Inputs
        self.inp_template = self.create_file_input(left_layout, "Template Word:", False)
        self.inp_excel = self.create_file_input(left_layout, "Planilha Excel:", False)
        self.inp_out_word = self.create_file_input(left_layout, "Saída Word (.docx):", True)
        self.inp_out_pdf = self.create_file_input(left_layout, "Saída PDF (.pdf):", True)

        left_layout.addStretch()

        # Botões de Ação
        btn_layout = QVBoxLayout()
        
        self.btn_full = QPushButton("▶️  (Opção 1) Gerar Relatório COMPLETO")
        self.btn_full.setStyleSheet("background-color: #0088cc; color: white;")
        self.btn_full.setCursor(Qt.PointingHandCursor)
        self.btn_full.clicked.connect(lambda: self.start_process(mode=1))
        
        self.btn_test = QPushButton("🧪  (Opção 2) Validar Formatação (Teste)")
        self.btn_test.setStyleSheet("background-color: #6c757d; color: white;")
        self.btn_test.setCursor(Qt.PointingHandCursor)
        self.btn_test.clicked.connect(lambda: self.start_process(mode=2))

        btn_layout.addWidget(self.btn_full)
        btn_layout.addWidget(self.btn_test)
        left_layout.addLayout(btn_layout)

        # --- LADO DIREITO: TERMINAL (CLI) ---
        right_panel = QWidget()
        right_layout = QVBoxLayout(right_panel)
        
        lbl_console = QLabel("🖥️ Terminal de Execução")
        lbl_console.setStyleSheet("font-size: 16px; font-weight: bold; color: #555;")
        right_layout.addWidget(lbl_console)

        self.console = RichConsoleWidget()
        right_layout.addWidget(self.console)

        # Adiciona ao layout principal
        main_layout.addWidget(left_panel)
        main_layout.addWidget(right_panel)

        # Define caminhos padrão
        self.set_defaults()

    def create_file_input(self, layout, label, is_save):
        lbl = QLabel(label)
        layout.addWidget(lbl)
        
        h = QHBoxLayout()
        inp = QLineEdit()
        btn = QPushButton("...")
        btn.setFixedWidth(40)
        btn.setStyleSheet("background-color: #eee; color: #333;")
        
        if is_save:
            btn.clicked.connect(lambda: self.browse_save(inp))
        else:
            btn.clicked.connect(lambda: self.browse_open(inp))
            
        h.addWidget(inp)
        h.addWidget(btn)
        layout.addLayout(h)
        return inp

    def set_defaults(self):
        # Ajuste aqui seus caminhos padrão
        p1 = r"C:\Users\pedrovictor.veras\OneDrive - Operador Nacional do Sistema Eletrico\Documentos\ESTAGIO_ONS_PVRV_2025\GitHub\Palkia-PDF-extractor\src\BulbassaurQT6-ETL\Perdas-Duplas-ETL-Desktop-V4\app\assets\docs\Lista de Contingências Duplas - Copia.docx"
        p2 = r"C:\Users\pedrovictor.veras\OneDrive - Operador Nacional do Sistema Eletrico\Documentos\ESTAGIO_ONS_PVRV_2025\GitHub\Palkia-PDF-extractor\src\BulbassaurQT6-ETL\Perdas-Duplas-ETL-Desktop-V4\app\assets\planilhas_PLC\perdas_duplas_ETL_corrigido.xlsx"
        
        self.inp_template.setText(p1)
        self.inp_excel.setText(p2)
        self.inp_out_word.setText("relatorio_final.docx")
        self.inp_out_pdf.setText("relatorio_final.pdf")

    # --- FUNÇÕES DE ARQUIVO ---
    def browse_open(self, line_edit):
        f, _ = QFileDialog.getOpenFileName(self, "Abrir Arquivo")
        if f: line_edit.setText(f)

    def browse_save(self, line_edit):
        f, _ = QFileDialog.getSaveFileName(self, "Salvar Arquivo")
        if f: line_edit.setText(f)

    # --- FUNÇÕES DE LOG E EXECUÇÃO ---
    def on_print(self, text):
        # Chamado toda vez que algo é printado no Python
        self.console.append_colored(text)

    def start_process(self, mode):
        t = self.inp_template.text()
        e = self.inp_excel.text()
        ow = self.inp_out_word.text()
        op = self.inp_out_pdf.text()

        # Feedback inicial
        self.console.clear()
        print("="*40)
        if mode == 1:
            print("🚀 INICIANDO PROCESSO COMPLETO (Excel -> Word -> PDF)")
        else:
            print("🧪 INICIANDO VALIDAÇÃO DE TABELA (Apenas teste)")
        print("="*40)

        # Trava botões
        self.btn_full.setEnabled(False)
        self.btn_test.setEnabled(False)

        # Inicia Thread
        self.worker = WorkerThread(t, e, ow, op, mode)
        self.worker.finished_signal.connect(self.on_finished)
        self.worker.error_signal.connect(self.on_error)
        self.worker.start()

    def on_finished(self):
        self.btn_full.setEnabled(True)
        self.btn_test.setEnabled(True)
        print("\n" + "="*40)
        print("✅ PROCESSO FINALIZADO COM SUCESSO")
        print("="*40)
        QMessageBox.information(self, "Concluído", "Operação finalizada! Verifique o log.")

    def on_error(self, err):
        self.btn_full.setEnabled(True)
        self.btn_test.setEnabled(True)
        print("\n" + "="*40)
        print(f"❌ ERRO FATAL: {err}")
        print("="*40)
        QMessageBox.critical(self, "Erro", f"Ocorreu um erro: {err}")

    def closeEvent(self, event):
        sys.stdout = self.sys_stdout # Restaura terminal original
        event.accept()

if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = MainWindow()
    window.show()
    sys.exit(app.exec())