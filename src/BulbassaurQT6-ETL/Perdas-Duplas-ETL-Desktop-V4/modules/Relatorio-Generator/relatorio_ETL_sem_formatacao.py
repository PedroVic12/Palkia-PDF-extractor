import sys
import os
import pandas as pd
from docx import Document
from docx.shared import RGBColor
from docx2pdf import convert
from docx.oxml.ns import qn

# --- IMPORTS PYSIDE6 ---
from PySide6.QtWidgets import (QWidget, QVBoxLayout, QHBoxLayout, QLabel, QLineEdit, 
                               QPushButton, QTextEdit, QFileDialog, QMessageBox, 
                               QFrame, QSpinBox, QProgressBar, QGroupBox)
from PySide6.QtCore import Qt, QThread, Signal, QObject

# ============================================================================
# 1. LÓGICA DE ANÁLISE E PROCESSAMENTO
# ============================================================================

def contar_paginas_quebras(doc_path):
    """
    Conta páginas baseando-se em Quebras de Página (Hard Breaks) explícitas.
    Retorna (num_paginas, lista_de_indices_paragrafos)
    """
    doc = Document(doc_path)
    paginas = 1 # Começa na página 1
    
    # Percorre todos os parágrafos procurando quebras
    for i, p in enumerate(doc.paragraphs):
        # Verifica quebras manuais (Ctrl+Enter)
        if 'lastRenderedPageBreak' in p._element.xml or '<w:br w:type="page"/>' in p._element.xml:
            paginas += 1
        
        # Verifica também nas Runs (às vezes a quebra fica dentro de um Run)
        for run in p.runs:
            if 'w:br' in run._element.xml and 'type="page"' in run._element.xml:
                paginas += 1
                break # Conta apenas uma quebra por parágrafo para não duplicar
                
    return paginas

def inserir_tabela_na_pagina(doc, table, pagina_alvo):
    """
    Tenta inserir a tabela na página alvo procurando por quebras de página.
    """
    if pagina_alvo <= 1:
        # Insere no início do documento
        doc.paragraphs[0].insert_paragraph_before("")._p.addnext(table._tbl)
        print("   📍 Inserido no INÍCIO do documento.")
        return

    paginas_contadas = 1
    inserido = False

    # Itera sobre o corpo do documento (parágrafos e tabelas)
    # Nota: Iterar sobre doc.paragraphs pula tabelas existentes, mas serve para achar as quebras
    for i, p in enumerate(doc.paragraphs):
        # Detecta quebra de página
        tem_quebra = False
        if '<w:br w:type="page"/>' in p._element.xml:
            tem_quebra = True
        else:
            for run in p.runs:
                if 'w:br' in run._element.xml and 'type="page"' in run._element.xml:
                    tem_quebra = True
                    break
        
        if tem_quebra:
            paginas_contadas += 1
            # Se acabamos de passar para a página desejada
            if paginas_contadas == pagina_alvo:
                # Insere logo após este parágrafo que contem a quebra
                p.insert_paragraph_before("") # Cria um espaço
                p.insert_paragraph_before("")._p.addnext(table._tbl)
                inserido = True
                print(f"   📍 Inserido logo após a quebra da página {paginas_contadas - 1}.")
                break
    
    if not inserido:
        print("   ⚠️ Página alvo não encontrada (documento menor que o esperado). Adicionando ao final.")
        doc.add_page_break()
        doc.add_paragraph("Tabela inserida aqui (Fim do Arquivo).")
        doc.element.body.append(table._tbl)

# ============================================================================
# 2. THREADS (WORKERS)
# ============================================================================

class AnaliseWorker(QThread):
    finished = Signal(dict) # Retorna dados da análise
    error = Signal(str)

    def __init__(self, doc_path, excel_path, sheet_name):
        super().__init__()
        self.doc = doc_path
        self.excel = excel_path
        self.sheet = sheet_name

    def run(self):
        try:
            resultados = {}
            
            # 1. Analisa Excel
            if os.path.exists(self.excel):
                try:
                    df = pd.read_excel(self.excel, sheet_name=self.sheet)
                except:
                    df = pd.read_excel(self.excel) # Tenta a primeira aba
                
                resultados['rows'] = df.shape[0]
                resultados['cols'] = df.shape[1]
                resultados['col_names'] = list(df.columns)
            
            # 2. Analisa Word
            if os.path.exists(self.doc):
                pgs = contar_paginas_quebras(self.doc)
                resultados['pages'] = pgs
            
            self.finished.emit(resultados)
        except Exception as e:
            self.error.emit(str(e))

class GeracaoWorker(QThread):
    finished = Signal()
    log = Signal(str)
    error = Signal(str)

    def __init__(self, config):
        super().__init__()
        self.cfg = config

    def run(self):
        try:
            self.log.emit("⏳ Iniciando processamento...")
            
            # 1. Carregar Excel
            df = pd.read_excel(self.cfg['excel'], sheet_name=self.cfg['sheet'])
            rows, cols = df.shape
            self.log.emit(f"✅ Excel carregado: {rows} linhas.")

            # 2. Abrir Word
            doc = Document(self.cfg['template'])
            
            # 3. Criar Tabela
            table = doc.add_table(rows=rows+1, cols=cols)
            table.style = 'Table Grid'
            
            # Cabeçalho Azul ONS
            hdr = table.rows[0].cells
            for idx, col in enumerate(df.columns):
                hdr[idx].text = str(col)
                run = hdr[idx].paragraphs[0].runs[0] if hdr[idx].paragraphs[0].runs else hdr[idx].paragraphs[0].add_run(str(col))
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                from docx.oxml import parse_xml
                from docx.oxml.ns import nsdecls
                hdr[idx]._element.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="005CA9"/>'.format(nsdecls('w'))))

            # Preencher Dados
            for r in range(rows):
                for c in range(cols):
                    val = df.values[r, c]
                    table.cell(r+1, c).text = str(val) if pd.notna(val) else "-"
            
            # 4. Inserir na Página Escolhida
            target = self.cfg['target_page']
            self.log.emit(f"📍 Tentando inserir na página {target}...")
            inserir_tabela_na_pagina(doc, table, target)

            # 5. Salvar
            doc.save(self.cfg['out_word'])
            self.log.emit("💾 Word Salvo.")
            
            # 6. PDF
            self.log.emit("🔄 Gerando PDF...")
            convert(os.path.abspath(self.cfg['out_word']), os.path.abspath(self.cfg['out_pdf']))
            self.log.emit("✅ Processo Finalizado!")
            
            self.finished.emit()

        except Exception as e:
            self.error.emit(str(e))

# ============================================================================
# 3. INTERFACE (IFRAME)
# ============================================================================

class RelatorioAnaliseWidget(QWidget):
    def __init__(self):
        super().__init__()
        self.init_ui()

    def init_ui(self):
        layout = QHBoxLayout(self)
        
        # --- COLUNA ESQUERDA (INPUTS) ---
        left_frame = QFrame()
        left_frame.setStyleSheet(".QFrame {background: white; border-radius: 8px; border: 1px solid #ddd;}")
        left_frame.setFixedWidth(400)
        left_vbox = QVBoxLayout(left_frame)
        
        left_vbox.addWidget(QLabel("📂 <b>Configuração de Arquivos</b>"))
        
        self.inp_template = self.create_file_input(left_vbox, "Template Word:", False)
        self.inp_excel = self.create_file_input(left_vbox, "Excel Dados:", False)
        
        left_vbox.addWidget(QLabel("Nome da Aba (Sheet):"))
        self.inp_sheet = QLineEdit("Sheet1")
        left_vbox.addWidget(self.inp_sheet)

        # Botão de Análise
        self.btn_analisar = QPushButton("🔍 1. ANALISAR ARQUIVOS")
        self.btn_analisar.setStyleSheet("background-color: #f39c12; color: white; font-weight: bold; padding: 10px;")
        self.btn_analisar.clicked.connect(self.analisar_arquivos)
        left_vbox.addWidget(self.btn_analisar)

        left_vbox.addSpacing(20)
        left_vbox.addWidget(QLabel("💾 <b>Configuração de Saída</b>"))
        self.inp_out_word = self.create_file_input(left_vbox, "Salvar Word:", True)
        self.inp_out_pdf = self.create_file_input(left_vbox, "Salvar PDF:", True)

        left_vbox.addStretch()
        layout.addWidget(left_frame)

        # --- COLUNA DIREITA (RESULTADOS + AÇÃO) ---
        right_frame = QFrame()
        right_vbox = QVBoxLayout(right_frame)

        # Grupo de Status
        gb_info = QGroupBox("📊 Estatísticas dos Arquivos")
        gb_info.setStyleSheet("QGroupBox {font-weight: bold; border: 1px solid #ccc; margin-top: 10px;} QGroupBox::title {subcontrol-origin: margin; left: 10px; padding: 0 5px;}")
        info_layout = QVBoxLayout(gb_info)
        
        self.lbl_excel_stats = QLabel("Excel: <i>Aguardando análise...</i>")
        self.lbl_word_stats = QLabel("Word: <i>Aguardando análise...</i>")
        self.lbl_word_stats.setStyleSheet("color: #666;")
        
        info_layout.addWidget(self.lbl_excel_stats)
        info_layout.addWidget(self.lbl_word_stats)
        right_vbox.addWidget(gb_info)

        # Seletor de Página
        right_vbox.addSpacing(10)
        row_page = QHBoxLayout()
        row_page.addWidget(QLabel("📍 <b>Inserir Tabela na Página:</b>"))
        self.spin_page = QSpinBox()
        self.spin_page.setRange(1, 999)
        self.spin_page.setValue(1)
        self.spin_page.setEnabled(False) # Habilita só após análise
        self.spin_page.setStyleSheet("font-size: 16px; padding: 5px;")
        row_page.addWidget(self.spin_page)
        right_vbox.addLayout(row_page)

        # Console Log
        self.console = QTextEdit()
        self.console.setReadOnly(True)
        self.console.setStyleSheet("background: #1e1e1e; color: #00ff00; font-family: Consolas;")
        right_vbox.addWidget(self.console)

        # Botão Final
        self.btn_gerar = QPushButton("🚀 2. GERAR RELATÓRIO")
        self.btn_gerar.setEnabled(False) # Habilita só após análise
        self.btn_gerar.setFixedHeight(50)
        self.btn_gerar.setStyleSheet("""
            QPushButton { background-color: #27ae60; color: white; font-weight: bold; font-size: 16px; border-radius: 5px; }
            QPushButton:disabled { background-color: #ccc; }
        """)
        self.btn_gerar.clicked.connect(self.gerar_relatorio)
        right_vbox.addWidget(self.btn_gerar)

        layout.addWidget(right_frame)

        self.set_defaults()

    def create_file_input(self, layout, label, save):
        layout.addWidget(QLabel(label))
        h = QHBoxLayout()
        le = QLineEdit()
        btn = QPushButton("...")
        btn.setFixedWidth(30)
        if save: btn.clicked.connect(lambda: self.browse_save(le))
        else: btn.clicked.connect(lambda: self.browse_open(le))
        h.addWidget(le)
        h.addWidget(btn)
        layout.addLayout(h)
        return le

    def browse_open(self, le):
        f, _ = QFileDialog.getOpenFileName(self, "Abrir Arquivo")
        if f: le.setText(f)

    def browse_save(self, le):
        f, _ = QFileDialog.getSaveFileName(self, "Salvar Arquivo")
        if f: le.setText(f)

    def set_defaults(self):
        # Ajuste seus caminhos aqui
        base = r"C:\Users\pedrovictor.veras\...\assets" # Coloque seu caminho base
        self.inp_out_word.setText("relatorio_final.docx")
        self.inp_out_pdf.setText("relatorio_final.pdf")

    # --- AÇÕES ---

    def analisar_arquivos(self):
        self.btn_analisar.setEnabled(False)
        self.btn_analisar.setText("⏳ Analisando...")
        self.console.append(">> Iniciando análise dos arquivos...")
        
        self.worker_analise = AnaliseWorker(
            self.inp_template.text(), 
            self.inp_excel.text(), 
            self.inp_sheet.text()
        )
        self.worker_analise.finished.connect(self.pos_analise)
        self.worker_analise.error.connect(lambda e: QMessageBox.critical(self, "Erro", e))
        self.worker_analise.start()

    def pos_analise(self, dados):
        self.btn_analisar.setEnabled(True)
        self.btn_analisar.setText("🔍 1. ANALISAR ARQUIVOS")
        
        # Atualiza labels
        self.lbl_excel_stats.setText(f"✅ Excel: <b>{dados.get('rows', '?')} linhas</b> x {dados.get('cols', '?')} colunas")
        self.lbl_excel_stats.setStyleSheet("color: green;")
        
        pgs = dados.get('pages', 1)
        self.lbl_word_stats.setText(f"✅ Word (Estimado): <b>{pgs} páginas</b> (Baseado em quebras)")
        self.lbl_word_stats.setStyleSheet("color: blue;")
        
        # Configura SpinBox
        self.spin_page.setEnabled(True)
        self.spin_page.setMaximum(pgs + 1) # Permite adicionar uma página extra no fim
        
        # Habilita botão final
        self.btn_gerar.setEnabled(True)
        self.console.append(f">> Análise concluída. O Word parece ter {pgs} seções/paginas.")
        self.console.append(">> Escolha a página e clique em Gerar.")

    def gerar_relatorio(self):
        cfg = {
            'template': self.inp_template.text(),
            'excel': self.inp_excel.text(),
            'sheet': self.inp_sheet.text(),
            'out_word': self.inp_out_word.text(),
            'out_pdf': self.inp_out_pdf.text(),
            'target_page': self.spin_page.value()
        }
        
        self.btn_gerar.setEnabled(False)
        self.worker_geracao = GeracaoWorker(cfg)
        self.worker_geracao.log.connect(self.console.append)
        self.worker_geracao.finished.connect(lambda: QMessageBox.information(self, "Sucesso", "Relatório Gerado!"))
        self.worker_geracao.finished.connect(lambda: self.btn_gerar.setEnabled(True))
        self.worker_geracao.error.connect(lambda e: QMessageBox.critical(self, "Erro", e))
        self.worker_geracao.start()

# --- PARA RODAR SEPARADO ---
if __name__ == "__main__":
    from PySide6.QtWidgets import QApplication
    app = QApplication(sys.argv)
    win = RelatorioAnaliseWidget()
    win.show()
    sys.exit(app.exec())