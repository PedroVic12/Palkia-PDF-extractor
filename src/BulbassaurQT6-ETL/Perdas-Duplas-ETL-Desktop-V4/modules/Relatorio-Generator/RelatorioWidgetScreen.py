import sys
import os
import pandas as pd
from docx import Document
from docx.shared import RGBColor
from docx2pdf import convert

# --- IMPORTS PYSIDE6 ---
from PySide6.QtWidgets import (QApplication, QWidget, QVBoxLayout, QHBoxLayout, QLabel, QLineEdit, 
                               QPushButton, QTextEdit, QFileDialog, QMessageBox, 
                               QFrame, QSpinBox, QGroupBox, QTabWidget, QTableWidget, 
                               QTableWidgetItem, QHeaderView, QSplitter)
from PySide6.QtCore import Qt, QThread, Signal, QObject, QUrl
from PySide6.QtGui import QFont, QColor

# Tenta importar o Visualizador Web (Necessário para ver o PDF dentro do App)
try:
    from PySide6.QtWebEngineWidgets import QWebEngineView, QWebEngineSettings
    HAS_WEBENGINE = True
except ImportError:
    HAS_WEBENGINE = False
    print("⚠️ Aviso: QtWebEngineWidgets não encontrado. A visualização de PDF interna será desativada.")

# ============================================================================
# 1. BACKEND (WORKERS & LÓGICA)
# ============================================================================

# ... (Lógica de Contar Páginas e Inserir Tabela mantida igual ao anterior) ...
def contar_paginas_quebras(doc_path):
    doc = Document(doc_path)
    paginas = 1
    for p in doc.paragraphs:
        if 'lastRenderedPageBreak' in p._element.xml or '<w:br w:type="page"/>' in p._element.xml:
            paginas += 1
        for run in p.runs:
            if 'w:br' in run._element.xml and 'type="page"' in run._element.xml:
                paginas += 1
                break
    return paginas

def inserir_tabela_na_pagina(doc, table, pagina_alvo):
    # (Mesma lógica robusta do código anterior)
    if pagina_alvo <= 1:
        doc.paragraphs[0].insert_paragraph_before("")._p.addnext(table._tbl)
        return

    paginas_contadas = 1
    inserido = False
    for p in doc.paragraphs:
        tem_quebra = False
        if '<w:br w:type="page"/>' in p._element.xml: tem_quebra = True
        else:
            for run in p.runs:
                if 'w:br' in run._element.xml and 'type="page"' in run._element.xml:
                    tem_quebra = True; break
        
        if tem_quebra:
            paginas_contadas += 1
            if paginas_contadas == pagina_alvo:
                p.insert_paragraph_before("") 
                p.insert_paragraph_before("")._p.addnext(table._tbl)
                inserido = True
                break
    
    if not inserido:
        doc.add_page_break()
        doc.element.body.append(table._tbl)

# --- WORKERS ---

class AnaliseWorker(QThread):
    finished = Signal(dict)
    error = Signal(str)

    def __init__(self, doc, excel, sheet):
        super().__init__()
        self.doc, self.excel, self.sheet = doc, excel, sheet

    def run(self):
        try:
            res = {}
            # Lê Excel completo para mostrar na aba de prévia
            if os.path.exists(self.excel):
                try: df = pd.read_excel(self.excel, sheet_name=self.sheet)
                except: df = pd.read_excel(self.excel)
                
                # Converte NaN para string vazia para exibição
                df = df.fillna("")
                res['df'] = df # Passamos o DataFrame inteiro
                res['rows'], res['cols'] = df.shape
            
            if os.path.exists(self.doc):
                res['pages'] = contar_paginas_quebras(self.doc)
            
            self.finished.emit(res)
        except Exception as e:
            self.error.emit(str(e))

class GeracaoWorker(QThread):
    finished = Signal()
    log = Signal(str)
    error = Signal(str)

    def __init__(self, config):
        super().__init__()
        self.cfg = config

    def run(self):
        try:
            self.log.emit("⏳ Carregando dados...")
            df = pd.read_excel(self.cfg['excel'], sheet_name=self.cfg['sheet'])
            doc = Document(self.cfg['template'])
            
            # Cria Tabela
            table = doc.add_table(rows=df.shape[0]+1, cols=df.shape[1])
            table.style = 'Table Grid'
            
            # Estilização ONS
            hdr = table.rows[0].cells
            for idx, col in enumerate(df.columns):
                hdr[idx].text = str(col)
                # (Lógica de XML para cor azul mantida)
                run = hdr[idx].paragraphs[0].runs[0] if hdr[idx].paragraphs[0].runs else hdr[idx].paragraphs[0].add_run(str(col))
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                from docx.oxml import parse_xml
                from docx.oxml.ns import nsdecls
                hdr[idx]._element.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="005CA9"/>'.format(nsdecls('w'))))

            # Dados
            self.log.emit("📝 Preenchendo tabela no Word...")
            for r in range(df.shape[0]):
                for c in range(df.shape[1]):
                    val = df.values[r, c]
                    table.cell(r+1, c).text = str(val) if pd.notna(val) else "-"
            
            inserir_tabela_na_pagina(doc, table, self.cfg['target_page'])
            
            doc.save(self.cfg['out_word'])
            self.log.emit("🔄 Convertendo para PDF...")
            convert(os.path.abspath(self.cfg['out_word']), os.path.abspath(self.cfg['out_pdf']))
            self.log.emit("✅ Finalizado!")
            self.finished.emit()

        except Exception as e:
            self.error.emit(str(e))

# ============================================================================
# 2. COMPONENTES VISUAIS (TABLE PREVIEW & PDF VIEWER)
# ============================================================================

class ExcelPreviewTable(QTableWidget):
    def __init__(self):
        super().__init__()
        self.setAlternatingRowColors(True)
        self.setStyleSheet("""
            QTableWidget { gridline-color: #ccc; font-family: 'Segoe UI'; font-size: 13px; }
            QHeaderView::section { background-color: #005CA9; color: white; font-weight: bold; padding: 5px; border: 1px solid #00407a; }
            QTableWidget::item { padding: 5px; }
        """)
    
    def load_dataframe(self, df):
        self.setRowCount(df.shape[0])
        self.setColumnCount(df.shape[1])
        self.setHorizontalHeaderLabels([str(c) for c in df.columns])
        
        for r in range(df.shape[0]):
            for c in range(df.shape[1]):
                val = str(df.iloc[r, c])
                item = QTableWidgetItem(val)
                item.setFlags(Qt.ItemIsEnabled | Qt.ItemIsSelectable) # Read-only
                self.setItem(r, c, item)
        
        self.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)

class PDFViewer(QWidget):
    def __init__(self):
        super().__init__()
        layout = QVBoxLayout(self)
        layout.setContentsMargins(0,0,0,0)
        
        if HAS_WEBENGINE:
            self.web_view = QWebEngineView()
            self.web_view.settings().setAttribute(QWebEngineSettings.PluginsEnabled, True)
            self.web_view.settings().setAttribute(QWebEngineSettings.PdfViewerEnabled, True)
            layout.addWidget(self.web_view)
        else:
            lbl = QLabel("Visualizador PDF não disponível.\n(Instale 'pip install PySide6-WebEngine' ou verifique sua instalação)")
            lbl.setAlignment(Qt.AlignCenter)
            layout.addWidget(lbl)
            self.web_view = None

    def load_pdf(self, file_path):
        if self.web_view and os.path.exists(file_path):
            # QWebEngine precisa de URL com file:///
            local_url = QUrl.fromLocalFile(os.path.abspath(file_path))
            self.web_view.setUrl(local_url)

# ============================================================================
# 3. WIDGET PRINCIPAL (COM ABAS)
# ============================================================================

class EmittingStream(QObject):
    textWritten = Signal(str)
    def write(self, text): self.textWritten.emit(str(text))
    def flush(self): pass

class RelatorioTabsWidget(QWidget):
    def __init__(self):
        super().__init__()
        self.init_ui()
        
        # Redireciona stdout
        self.sys_stdout = sys.stdout
        self.emitter = EmittingStream()
        self.emitter.textWritten.connect(self.console.append)
        sys.stdout = self.emitter

    def init_ui(self):
        layout = QVBoxLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)

        # --- A. ABAS PRINCIPAIS ---
        self.tabs = QTabWidget()
        self.tabs.setStyleSheet("""
            QTabWidget::pane { border: 1px solid #ccc; background: white; }
            QTabBar::tab { background: #f0f0f0; padding: 10px 20px; font-weight: bold; color: #555; }
            QTabBar::tab:selected { background: #005CA9; color: white; }
        """)

        # Aba 1: Painel de Controle
        self.tab_controle = QWidget()
        self.setup_tab_controle()
        
        # Aba 2: Prévia Excel
        self.tab_excel = ExcelPreviewTable()
        
        # Aba 3: Visualizador PDF
        self.tab_pdf = PDFViewer()

        self.tabs.addTab(self.tab_controle, "🎛️ Painel de Controle")
        self.tabs.addTab(self.tab_excel, "📊 Prévia dos Dados (Excel)")
        self.tabs.addTab(self.tab_pdf, "📄 Visualizador do Relatório (PDF)")

        layout.addWidget(self.tabs)

        self.set_defaults()

    def setup_tab_controle(self):
        layout = QHBoxLayout(self.tab_controle)
        
        # Coluna Esquerda: Configurações
        left_frame = QFrame()
        left_frame.setFixedWidth(400)
        vbox = QVBoxLayout(left_frame)
        
        vbox.addWidget(QLabel("📂 <b>Entradas:</b>"))
        self.inp_template = self.create_file_input(vbox, "Template Word:", False)
        self.inp_excel = self.create_file_input(vbox, "Excel Dados:", False)
        
        hbox_sheet = QHBoxLayout()
        hbox_sheet.addWidget(QLabel("Aba do Excel:"))
        self.inp_sheet = QLineEdit("Dados")
        hbox_sheet.addWidget(self.inp_sheet)
        vbox.addLayout(hbox_sheet)

        vbox.addSpacing(10)
        vbox.addWidget(QLabel("💾 <b>Saídas:</b>"))
        self.inp_out_word = self.create_file_input(vbox, "Salvar Word:", True)
        self.inp_out_pdf = self.create_file_input(vbox, "Salvar PDF:", True)
        
        vbox.addSpacing(20)
        
        # Botão Analisar
        self.btn_analisar = QPushButton("🔍 1. ANALISAR & PRÉ-VISUALIZAR")
        self.btn_analisar.setFixedHeight(45)
        self.btn_analisar.setStyleSheet("background-color: #f39c12; color: white; font-weight: bold; border-radius: 5px;")
        self.btn_analisar.clicked.connect(self.analisar)
        vbox.addWidget(self.btn_analisar)

        # Status Análise
        self.lbl_status_analise = QLabel("<i>Aguardando análise...</i>")
        self.lbl_status_analise.setStyleSheet("color: #666; margin-top: 5px;")
        self.lbl_status_analise.setWordWrap(True)
        vbox.addWidget(self.lbl_status_analise)

        # Configuração Pagina (Desabilitado até analisar)
        vbox.addSpacing(10)
        self.group_page = QGroupBox("Onde inserir?")
        self.group_page.setEnabled(False)
        l_page = QHBoxLayout(self.group_page)
        l_page.addWidget(QLabel("Inserir na Página:"))
        self.spin_page = QSpinBox()
        self.spin_page.setRange(1, 999)
        l_page.addWidget(self.spin_page)
        vbox.addWidget(self.group_page)

        vbox.addStretch()

        # Botão Gerar
        self.btn_gerar = QPushButton("🚀 2. GERAR E VER PDF")
        self.btn_gerar.setEnabled(False)
        self.btn_gerar.setFixedHeight(50)
        self.btn_gerar.setStyleSheet("""
            QPushButton { background-color: #27ae60; color: white; font-weight: bold; border-radius: 5px; }
            QPushButton:disabled { background-color: #ccc; }
        """)
        self.btn_gerar.clicked.connect(self.gerar)
        vbox.addWidget(self.btn_gerar)

        layout.addWidget(left_frame)

        # Coluna Direita: Console
        right_vbox = QVBoxLayout()
        lbl_term = QLabel("🖥️ Terminal de Processamento")
        lbl_term.setStyleSheet("font-weight: bold;")
        right_vbox.addWidget(lbl_term)
        
        self.console = QTextEdit()
        self.console.setReadOnly(True)
        self.console.setStyleSheet("background-color: #1e1e1e; color: #00ff00; font-family: Consolas;")
        right_vbox.addWidget(self.console)
        
        layout.addLayout(right_vbox)

    def create_file_input(self, layout, label, save):
        l = QLabel(label)
        layout.addWidget(l)
        h = QHBoxLayout()
        le = QLineEdit()
        btn = QPushButton("...")
        btn.setFixedWidth(30)
        if save: btn.clicked.connect(lambda: self.browse_save(le))
        else: btn.clicked.connect(lambda: self.browse_open(le))
        h.addWidget(le)
        h.addWidget(btn)
        layout.addLayout(h)
        return le

    # --- Lógica de Ação ---

    def analisar(self):
        self.btn_analisar.setEnabled(False)
        self.console.clear()
        self.console.append(">> Iniciando análise...")
        
        self.worker_analise = AnaliseWorker(self.inp_template.text(), self.inp_excel.text(), self.inp_sheet.text())
        self.worker_analise.finished.connect(self.pos_analise)
        self.worker_analise.error.connect(lambda e: QMessageBox.critical(self, "Erro", e))
        self.worker_analise.start()

    def pos_analise(self, dados):
        self.btn_analisar.setEnabled(True)
        
        # 1. Popula Aba Excel
        if 'df' in dados:
            self.tab_excel.load_dataframe(dados['df'])
            self.lbl_status_analise.setText(f"✅ Excel OK: {dados['rows']} linhas.\n✅ Word OK: {dados.get('pages', '?')} págs.")
            # Habilita aba excel e pisca ela
            self.tabs.setTabEnabled(1, True)
            
            # 2. Configura Paginação
            if 'pages' in dados:
                self.spin_page.setMaximum(dados['pages'] + 1)
                self.group_page.setEnabled(True)
                self.btn_gerar.setEnabled(True)
            
            self.console.append(">> Análise concluída. Veja a aba 'Prévia dos Dados' para conferir a tabela.")
            QMessageBox.information(self, "Análise", "Arquivos analisados!\nVerifique a aba 'Prévia dos Dados'.")

    def gerar(self):
        self.btn_gerar.setEnabled(False)
        self.console.append(">> Iniciando geração do relatório...")
        
        cfg = {
            'template': self.inp_template.text(),
            'excel': self.inp_excel.text(),
            'sheet': self.inp_sheet.text(),
            'out_word': self.inp_out_word.text(),
            'out_pdf': self.inp_out_pdf.text(),
            'target_page': self.spin_page.value()
        }
        
        self.worker_geracao = GeracaoWorker(cfg)
        self.worker_geracao.log.connect(self.console.append)
        self.worker_geracao.finished.connect(self.pos_geracao)
        self.worker_geracao.error.connect(lambda e: QMessageBox.critical(self, "Erro", e))
        self.worker_geracao.start()

    def pos_geracao(self):
        self.btn_gerar.setEnabled(True)
        
        # Carrega o PDF na Aba 3
        pdf_path = self.inp_out_pdf.text()
        self.tab_pdf.load_pdf(pdf_path)
        
        self.console.append(">> PDF Carregado no Visualizador.")
        
        # Muda o foco para a aba do PDF
        self.tabs.setCurrentIndex(2) 
        QMessageBox.information(self, "Sucesso", "Relatório Gerado e Aberto na Visualização!")

    # Funções auxiliares de arquivo
    def browse_open(self, le):
        f, _ = QFileDialog.getOpenFileName(self, "Abrir")
        if f: le.setText(f)
    def browse_save(self, le):
        f, _ = QFileDialog.getSaveFileName(self, "Salvar")
        if f: le.setText(f)
    
    def set_defaults(self):
        # Defina seus caminhos padrão aqui
        pass

    def closeEvent(self, event):
        sys.stdout = self.sys_stdout
        event.accept()

# --- MAIN PARA TESTE ISOLADO ---
if __name__ == "__main__":
    app = QApplication(sys.argv)
    app.setStyle("Fusion")
    app.setFont(QFont("Segoe UI", 10))
    win = RelatorioTabsWidget()
    win.setWindowTitle("Sistema Desktop - Gerador de Relatórios de Perdas Duplas ETL")
    win.resize(1200, 800)
    win.show()
    sys.exit(app.exec())