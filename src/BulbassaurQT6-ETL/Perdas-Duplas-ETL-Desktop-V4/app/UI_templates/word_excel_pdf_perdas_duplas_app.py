import sys
import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.shared import Inches
import os
import time
from docx2pdf import convert
from PySide6.QtWidgets import (QApplication, QMainWindow, QWidget, QVBoxLayout, 
                               QHBoxLayout, QLabel, QLineEdit, QPushButton, 
                               QProgressBar, QFileDialog, QMessageBox, QFrame)
from PySide6.QtCore import Qt, QThread, Signal

# --- Lógica de Automação (Executada em uma Thread Separada) ---

class WorkerThread(QThread):
    # Sinais personalizados para comunicação com a GUI
    progress_signal = Signal(int, str)
    finished_signal = Signal(str)
    error_signal = Signal(str)

    def __init__(self, template_path, excel_path, output_docx, output_pdf):
        super().__init__()
        self.template_path = template_path
        self.excel_path = excel_path
        self.output_docx = output_docx
        self.output_pdf = output_pdf

    def run(self):
        try:
            self.progress_signal.emit(10, "Verificando arquivos...")
            if not os.path.exists(self.template_path) or not os.path.exists(self.excel_path):
                raise FileNotFoundError("Template Word ou arquivo Excel não encontrado.")

            self.progress_signal.emit(20, "Lendo dados do Excel e Word...")
            df = pd.read_excel(self.excel_path)
            doc = Document(self.template_path)
            
            # Simula a contagem de páginas para o progresso
            num_pages_template = len(doc.sections) 
            self.progress_signal.emit(30, f"Template com {num_pages_template} seções.")

            self.progress_signal.emit(40, "Adicionando nova página e cabeçalho...")
            # Adiciona uma quebra de seção (nova página)
            doc.add_section(WD_SECTION.NEW_PAGE)
            doc.add_heading('Relatório Gerado Automaticamente', level=1)
            doc.add_paragraph("Dados inseridos via script Python/PySide6:")

            rows, cols = df.shape
            table = doc.add_table(rows + 1, cols)
            table.style = 'Table Grid'

            # Preenche o cabeçalho
            for j, col_name in enumerate(df.columns):
                table.cell(0, j).text = str(col_name)

            # Preenche os dados com progresso simulado (progresso de 40 a 70%)
            for i in range(rows):
                for j in range(cols):
                    table.cell(i + 1, j).text = str(df.values[i, j])
                
                # Atualiza o progresso a cada X linhas para não sobrecarregar a GUI
                if i % 50 == 0: 
                    progress_val = 40 + int(30 * (i / rows))
                    self.progress_signal.emit(progress_val, f"Preenchendo tabela ({i+1}/{rows} linhas)...")

            self.progress_signal.emit(70, "Salvando documento Word temporário...")
            doc.save(self.output_docx)
            time.sleep(1) # Pequena pausa para a GUI atualizar

            self.progress_signal.emit(80, "Iniciando conversão para PDF (Requer MS Office)...")
            # Esta é a parte que pode demorar ou falhar se o Office não estiver instalado
            convert(self.output_docx, self.output_pdf)
            
            self.progress_signal.emit(100, f"Concluído! PDF salvo em: {self.output_pdf}")
            self.finished_signal.emit(self.output_pdf)

        except Exception as e:
            self.error_signal.emit(str(e))


# --- Interface Gráfica PySide6 ---

class AutomationWidget(QWidget):
    def __init__(self):
        super().__init__()
        self.initUI()
        self.worker = None

    def initUI(self):
        # Layout Principal (Vertical)
        layout = QVBoxLayout()
        layout.setSpacing(15)

        # 1. Título
        title = QLabel("Automação Word/Excel/PDF")
        title.setStyleSheet("font-size: 18px; font-weight: bold; margin-bottom: 10px;")
        layout.addWidget(title)

        # 2. Input Template Word
        layout.addWidget(QLabel("Caminho do Template Word (.docx):"))
        self.entry_word = QLineEdit(r"C:\Users\pedrovictor.veras\OneDrive - Operador Nacional do Sistema Eletrico\Documentos\ESTAGIO_ONS_PVRV_2025\GitHub\Palkia-PDF-extractor\src\BulbassaurQT6-ETL\sistema-ferramentas-RPA-desktop\app\assets\docs\Lista de Contingências Duplas - Copia.docx")
        btn_browse_word = QPushButton("Navegar...")
        btn_browse_word.clicked.connect(lambda: self.browse_file(self.entry_word, "*.docx"))
        
        h_layout_word = QHBoxLayout()
        h_layout_word.addWidget(self.entry_word)
        h_layout_word.addWidget(btn_browse_word)
        layout.addLayout(h_layout_word)

        # 3. Input Excel Data
        layout.addWidget(QLabel("Caminho dos Dados Excel (.xlsx):"))
        self.entry_excel = QLineEdit(r"C:\Users\pedrovictor.veras\OneDrive - Operador Nacional do Sistema Eletrico\Documentos\ESTAGIO_ONS_PVRV_2025\GitHub\Palkia-PDF-extractor\src\BulbassaurQT6-ETL\sistema-ferramentas-RPA-desktop\app\assets\planilhas_PLC\perdas_duplas_ETL_corrigido.xlsx")
        btn_browse_excel = QPushButton("Navegar...")
        btn_browse_excel.clicked.connect(lambda: self.browse_file(self.entry_excel, "*.xlsx"))

        h_layout_excel = QHBoxLayout()
        h_layout_excel.addWidget(self.entry_excel)
        h_layout_excel.addWidget(btn_browse_excel)
        layout.addLayout(h_layout_excel)

        # Adiciona uma linha horizontal para separar
        line = QFrame()
        line.setFrameShape(QFrame.HLine)
        line.setFrameShadow(QFrame.Sunken)
        layout.addWidget(line)

        # 4. Botão de Execução
        self.btn_run = QPushButton("Executar Automação e Gerar PDF")
        self.btn_run.setStyleSheet("padding: 10px; background-color: #4CAF50; color: white; font-weight: bold;")
        self.btn_run.clicked.connect(self.start_automation)
        layout.addWidget(self.btn_run)

        # 5. Barra de Progresso
        self.progress_bar = QProgressBar()
        self.progress_bar.setValue(0)
        layout.addWidget(self.progress_bar)

        # 6. Status Label
        self.status_label = QLabel("Aguardando início...")
        self.status_label.setStyleSheet("color: gray; font-style: italic;")
        layout.addWidget(self.status_label)
        
        self.setLayout(layout)

    def browse_file(self, line_edit, filter_str):
        file_path, _ = QFileDialog.getOpenFileName(self, "Selecionar Arquivo", "", f"Arquivos ({filter_str});;Todos os Arquivos (*)")
        if file_path:
            line_edit.setText(file_path)

    def start_automation(self):
        template_path = self.entry_word.text()
        excel_path = self.entry_excel.text()
        output_docx = "documento_final_temp.docx"
        output_pdf = "relatorio_final.pdf" # Salva no diretório atual do script

        if not template_path or not excel_path:
            QMessageBox.warning(self, "Aviso", "Por favor, preencha ambos os caminhos de arquivo.")
            return

        self.btn_run.setEnabled(False)
        self.status_label.setText("Iniciando processamento em segundo plano...")
        self.progress_bar.setValue(0)

        # Inicia a lógica em uma thread separada
        self.worker = WorkerThread(template_path, excel_path, output_docx, output_pdf)
        self.worker.progress_signal.connect(self.update_progress)
        self.worker.finished_signal.connect(self.on_finished)
        self.worker.error_signal.connect(self.on_error)
        self.worker.start()

    def update_progress(self, value, message):
        self.progress_bar.setValue(value)
        self.status_label.setText(message)

    def on_finished(self, output_pdf_path):
        self.btn_run.setEnabled(True)
        self.status_label.setText(f"Processo concluído. PDF gerado: {output_pdf_path}")
        QMessageBox.information(self, "Sucesso", f"O arquivo PDF foi gerado com sucesso em:\n{output_pdf_path}")

    def on_error(self, error_message):
        self.btn_run.setEnabled(True)
        self.status_label.setText(f"ERRO: {error_message}")
        QMessageBox.critical(self, "Erro na Automação", f"Ocorreu um erro:\n{error_message}")


class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Sistema de Automação PLC")
        self.setGeometry(100, 100, 600, 450)
        
        # Centraliza o widget principal (nosso "iframe")
        self.central_widget = AutomationWidget()
        self.setCentralWidget(self.central_widget)

# --- Execução Principal da Aplicação ---
if __name__ == "__main__":
    # Garante que o MS Office esteja instalado, pois docx2pdf requer isso no Windows.
    # Não há como verificar isso em Python puro, então assumimos que está instalado.
    app = QApplication(sys.argv)
    mainWin = MainWindow()
    mainWin.show()
    sys.exit(app.exec())
