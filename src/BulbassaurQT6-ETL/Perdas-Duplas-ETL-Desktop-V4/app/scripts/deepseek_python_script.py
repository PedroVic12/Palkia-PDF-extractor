import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx2pdf import convert
import os
from tqdm import tqdm
import time

def adicionar_pagina_e_tabela(doc_path: str, excel_path: str, output_docx_path: str):
    print(f"Abrindo documento Word template: {doc_path}")
    doc = Document(doc_path)
    
    # Adiciona a nova seção e configura margens
    nova_secao = doc.add_section(WD_SECTION.NEW_PAGE)
    nova_secao.top_margin = Inches(0.5)
    nova_secao.bottom_margin = Inches(0.5)
    nova_secao.left_margin = Inches(0.5)
    nova_secao.right_margin = Inches(0.5)

    # Adiciona título centralizado
    titulo = doc.add_heading('Relatório de Perdas Duplas', level=1)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph("Abaixo estão os dados extraídos da planilha Excel:")

    # Carrega os dados do Excel
    df = pd.read_excel(excel_path)
    rows, cols = df.shape
    
    # Cria tabela com estilo profissional
    table = doc.add_table(rows + 1, cols)
    table.style = 'Light Grid Accent 1'  # Estilo mais profissional
    
    # Configura a tabela para ocupar toda a largura da página
    table.autofit = False
    table.allow_autofit = True
    
    # Largura total da página (margens de 0.5" cada lado = 7.5" disponíveis)
    total_width = Inches(7.5)
    
    # Calcula largura das colunas (distribuição igual ou personalizada)
    col_widths = [total_width / cols for _ in range(cols)]
    
    # Preenche o cabeçalho da tabela com formatação
    for j, column_name in enumerate(df.columns):
        # Configura largura da coluna
        table.columns[j].width = col_widths[j]
        
        # Formata célula do cabeçalho
        cell = table.cell(0, j)
        cell.text = str(column_name)
        
        # Formatação do cabeçalho
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(11)
                run.font.name = 'Arial'
        
        # Adiciona cor de fundo ao cabeçalho
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), 'D9D9D9')  # Cinza claro
        cell._element.get_or_add_tcPr().append(shading_elm)

    # Preenche os dados com barra de progresso e formatação
    print(f"Preenchendo tabela com {rows} linhas...")
    for i in tqdm(range(rows), desc="Progresso da Tabela"):
        for j in range(cols):
            cell = table.cell(i + 1, j)
            cell.text = str(df.values[i, j])
            
            # Formatação das células de dados
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    run.font.name = 'Arial'

    # Ajusta layout final da tabela
    for row in table.rows:
        row.height = Inches(0.3)  # Altura uniforme das linhas

    doc.save(output_docx_path)
    print(f"Documento Word modificado salvo em '{output_docx_path}'.")
    return output_docx_path

def converter_docx_para_pdf(docx_path: str, pdf_path: str):
    print(f"Iniciando conversão de DOCX para PDF. Isso pode demorar...")
    start_time = time.time()
    
    try:
        convert(docx_path, pdf_path)
        end_time = time.time()
        print(f"Conversão concluída em {end_time - start_time:.2f} segundos.")
        print(f"Arquivo PDF salvo em '{pdf_path}'.")
    except Exception as e:
        print(f"ERRO na conversão para PDF. Certifique-se que o Microsoft Office está instalado.")
        print(f"Detalhes do erro: {e}")
        raise

# --- MODO DEBUG ---
def modo_debug():
    """Modo para testar e ajustar a formatação"""
    print("=== MODO DEBUG ATIVADO ===")
    
    template_word_path = r"C:\Users\pedrovictor.veras\OneDrive - Operador Nacional do Sistema Eletrico\Documentos\ESTAGIO_ONS_PVRV_2025\GitHub\Palkia-PDF-extractor\src\BulbassaurQT6-ETL\sistema-ferramentas-RPA-desktop\app\assets\docs\Lista de Contingências Duplas - Copia.docx" 
    excel_data_path = r"C:\Users\pedrovictor.veras\OneDrive - Operador Nacional do Sistema Eletrico\Documentos\ESTAGIO_ONS_PVRV_2025\GitHub\Palkia-PDF-extractor\src\BulbassaurQT6-ETL\sistema-ferramentas-RPA-desktop\app\assets\planilhas_PLC\perdas_duplas_ETL_corrigido.xlsx"
    
    output_word_path = "documento_debug.docx"
    
    if os.path.exists(template_word_path) and os.path.exists(excel_data_path):
        try:
            # Gera apenas o DOCX para verificar a formatação
            docx_gerado_path = adicionar_pagina_e_tabela(
                template_word_path, 
                excel_data_path, 
                output_word_path
            )
            print(f"\n✅ Documento de debug gerado: {docx_gerado_path}")
            print("Abra o arquivo para verificar a formatação antes de gerar o PDF.")
            
        except Exception as e:
            print(f"\n❌ Erro no modo debug: {e}")
    else:
        print("\n❌ Arquivos de entrada não encontrados.")

# --- MODO DEPLOY ---
def modo_deploy():
    """Modo para gerar o PDF final"""
    print("=== MODO DEPLOY ATIVADO ===")
    
    template_word_path = r"C:\Users\pedrovictor.veras\OneDrive - Operador Nacional do Sistema Eletrico\Documentos\ESTAGIO_ONS_PVRV_2025\GitHub\Palkia-PDF-extractor\src\BulbassaurQT6-ETL\sistema-ferramentas-RPA-desktop\app\assets\docs\Lista de Contingências Duplas - Copia.docx" 
    excel_data_path = r"C:\Users\pedrovictor.veras\OneDrive - Operador Nacional do Sistema Eletrico\Documentos\ESTAGIO_ONS_PVRV_2025\GitHub\Palkia-PDF-extractor\src\BulbassaurQT6-ETL\sistema-ferramentas-RPA-desktop\app\assets\planilhas_PLC\perdas_duplas_ETL_corrigido.xlsx"
    
    output_word_path = "documento_final_gerado.docx"
    output_pdf_path = "relatorio_PLC_perdas_duplas.pdf"

    if os.path.exists(template_word_path) and os.path.exists(excel_data_path):
        try:
            # Etapa 1: Gerar DOCX
            docx_gerado_path = adicionar_pagina_e_tabela(
                template_word_path, 
                excel_data_path, 
                output_word_path
            )
            
            # Etapa 2: Converter para PDF
            converter_docx_para_pdf(docx_gerado_path, output_pdf_path)
            
            print("\n✅ PROCESSO CONCLUÍDO! PDF gerado com sucesso.")
            
        except Exception as e:
            print(f"\n❌ Erro no modo deploy: {e}")
    else:
        print("\n❌ Arquivos de entrada não encontrados.")

# --- Execução Principal ---
if __name__ == "__main__":
    print("Escolha o modo de execução:")
    print("1 - DEBUG (gera apenas DOCX para testar formatação)")
    print("2 - DEPLOY (gera DOCX e converte para PDF)")
    
    escolha = input("Digite 1 ou 2: ").strip()
    
    if escolha == "1":
        modo_debug()
    elif escolha == "2":
        modo_deploy()
    else:
        print("Opção inválida. Execute novamente.")