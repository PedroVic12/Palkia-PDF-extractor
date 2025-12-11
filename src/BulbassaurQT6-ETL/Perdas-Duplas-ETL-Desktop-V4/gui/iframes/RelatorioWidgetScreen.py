# RelatorioWidgetFrame.py
import sys
import os
import pandas as pd
from PIL import Image
import fitz  # PyMuPDF

from docx import Document
from docx.shared import RGBColor, Cm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx2pdf import convert

from PySide6.QtWidgets import (
    QApplication, QFrame, QWidget, QLabel, QVBoxLayout, QHBoxLayout,
    QPushButton, QFileDialog, QLineEdit, QTextEdit, QMessageBox,
    QGroupBox, QSpinBox, QTabWidget, QScrollArea, QTableWidget,
    QTableWidgetItem, QHeaderView, QFormLayout, QComboBox, QProgressBar
)
from PySide6.QtCore import Qt, QThread, Signal, Slot
from PySide6.QtGui import QPixmap, QFont
from PySide6.QtWebEngineWidgets import QWebEngineView

# =======================================================
#  UTILITÁRIOS DE FORMATAÇÃO DOCX
# =======================================================

def set_cell_shading(cell, fill_hex: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def set_run_color(run, hex_color: str):
    r = int(hex_color[0:2], 16)
    g = int(hex_color[2:4], 16)
    b = int(hex_color[4:6], 16)
    run.font.color.rgb = RGBColor(r, g, b)

def px_to_cm(px):
    return px / 37.8


def gerar_preview_html(df, widths=None, header_fill="4472C4", header_text="FFFFFF"):
    bg = f"#{header_fill}"
    color = f"#{header_text}"
    html = '<html><head><meta charset="utf-8"></head><body>'
    html += '<table border="1" cellspacing="0" cellpadding="6" style="border-collapse:collapse;font-family:Arial;width:100%;">'
    # cabeçalho
    html += '<tr>'
    for col in df.columns:
        w = widths.get(col, 140) if widths else 140
        html += f'<th style="background:{bg};color:{color};width:{w}px;">{col}</th>'
    html += '</tr>'
    # linhas
    for _, row in df.iterrows():
        html += '<tr>'
        for val in row:
            html += f'<td>{"" if pd.isna(val) else val}</td>'
        html += '</tr>'
    html += '</table>'
    html += '</body></html>'
    return html

# =======================================================
#  PREVIEW EXCEL (QTableWidget)
# =======================================================
class ExcelPreview(QTableWidget):
    def load_df(self, df):
        self.setRowCount(df.shape[0])
        self.setColumnCount(df.shape[1])
        self.setHorizontalHeaderLabels([str(c) for c in df.columns])

        for r in range(df.shape[0]):
            for c in range(df.shape[1]):
                item = QTableWidgetItem(str(df.iloc[r, c]))
                self.setItem(r, c, item)

        self.setAlternatingRowColors(True)
        self.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        self.verticalHeader().setVisible(False)

        self.setStyleSheet("""
            QTableWidget {
                font-size: 13px;
                gridline-color: #CCC;
            }
            QHeaderView::section {
                background: #005CA9;
                color: white;
                padding: 4px;
                font-weight: bold;
            }
        """)


# =======================================================
#  PREVIEW DOCUMENTO — CARROSSEL DE PÁGINAS
# =======================================================
class DocumentPreview(QWidget):
    def __init__(self):
        super().__init__()
        self.layout = QVBoxLayout(self)
        self.layout.setAlignment(Qt.AlignTop)

        self.scroll = QScrollArea()
        self.scroll.setWidgetResizable(True)

        self.container = QWidget()
        self.vbox = QVBoxLayout(self.container)
        self.vbox.setAlignment(Qt.AlignTop)

        self.scroll.setWidget(self.container)
        self.layout.addWidget(self.scroll)

    def load_pdf(self, path):
        # remove previews antigos
        self.clear_preview()

        if not os.path.exists(path):
            return

        pdf = fitz.open(path)
        for page in pdf:
            pix = page.get_pixmap(dpi=150)
            img_path = f"{path}_page_{page.number}.png"
            pix.save(img_path)

            lbl = QLabel()
            pixmap = QPixmap(img_path)
            lbl.setPixmap(pixmap)
            lbl.setScaledContents(True)
            lbl.setMaximumWidth(900)

            self.vbox.addWidget(lbl)

    def load_word(self, path):
        # render Word -> PDF -> preview
        temp_pdf = path.replace(".docx", "_preview.pdf")
        os.system(f'libreoffice --headless --convert-to pdf "{path}" --outdir "{os.path.dirname(path)}"')
        self.load_pdf(temp_pdf)

    def clear_preview(self):
        while self.vbox.count():
            w = self.vbox.takeAt(0).widget()
            if w:
                w.deleteLater()

# -----------------------------
# Função para listar estilos de tabela do Word (a partir do template)
# -----------------------------
def listar_estilos_de_tabela(template_path):
    try:
        doc = Document(template_path)
        estilos = [s.name for s in doc.styles]
        # garantir alguns estilos padrão caso o template não traga
        defaults = ["Table Grid", "Light Shading", "Medium Shading 1", "List Table 1 - Accent 1"]
        for d in defaults:
            if d not in estilos:
                estilos.append(d)
        return estilos
    except Exception:
        return ["Table Grid"]

# -----------------------------
# Worker: gera o documento em background
# -----------------------------
class Worker(QThread):
    progresso = Signal(int)
    mensagem = Signal(str)
    completo = Signal(str)

    def __init__(self, excel_path, template_path, pagina_destino, incluir_teste, converter_pdf, col_widths_map, table_style):
        super().__init__()
        self.excel_path = excel_path
        self.template_path = template_path
        self.pagina_destino = pagina_destino
        self.incluir_teste = incluir_teste
        self.converter_pdf = converter_pdf
        self.col_widths_map = col_widths_map  # dict col -> px
        self.table_style = table_style

    def run(self):
        try:
            # Ler excel
            self.mensagem.emit("Carregando planilha Excel...")
            df = pd.read_excel(self.excel_path)

            self.mensagem.emit("Abrindo template Word...")
            doc = Document(self.template_path)

            # Tabela de teste (opcional)
            if self.incluir_teste:
                self.mensagem.emit("Criando tabela de teste...")
                tabela_test = doc.add_table(rows=4, cols=3)
                tabela_test.style = self.table_style if self.table_style else "Table Grid"
                header_fill = "4472C4"
                header_text_color = "FFFFFF"
                for i in range(3):
                    cell = tabela_test.cell(0, i)
                    run = cell.paragraphs[0].add_run(f"CABEÇALHO {i+1}")
                    run.bold = True
                    set_cell_shading(cell, header_fill)
                    set_run_color(run, header_text_color)
                for i in range(1, 4):
                    for j in range(3):
                        tabela_test.cell(i, j).text = f"Teste {i}-{j+1}"
                doc.add_page_break()

            # Criar tabela principal
            self.mensagem.emit("Criando tabela principal...")
            rows = df.shape[0] + 1
            cols = df.shape[1] if df.shape[1] > 0 else 1
            tabela = doc.add_table(rows=rows, cols=cols)
            tabela.style = self.table_style if self.table_style else "Table Grid"

            # Cabeçalho
            header_fill = "4472C4"
            header_text_color = "FFFFFF"
            for j, col in enumerate(df.columns):
                cell = tabela.cell(0, j)
                run = cell.paragraphs[0].add_run(str(col))
                run.bold = True
                set_cell_shading(cell, header_fill)
                set_run_color(run, header_text_color)

            # Dados e progresso
            total = df.shape[0] if df.shape[0] > 0 else 1
            for i, row in enumerate(df.values):
                for j, val in enumerate(row):
                    tabela.cell(i+1, j).text = "" if pd.isna(val) else str(val)
                if (i+1) % 1 == 0:
                    pct = int(((i+1) / total) * 100)
                    self.progresso.emit(pct)

            # Aplicar larguras por coluna (converter px -> cm)
            self.mensagem.emit("Aplicando larguras de coluna...")
            try:
                for j, col in enumerate(df.columns):
                    px = int(self.col_widths_map.get(col, 140))
                    cm_val = px_to_cm(px)
                    # set width for each cell in column
                    for r in range(rows):
                        tabela.cell(r, j).width = Cm(cm_val)
                    # also try to set tabela.columns[j].width if exists
                    try:
                        tabela.columns[j].width = Cm(cm_val)
                    except Exception:
                        pass
            except Exception as e:
                # não fatal: reporta e continua
                self.mensagem.emit(f"Aviso: não foi possível aplicar todas as larguras ({e})")

            # Inserir tabela na página destino (simples, coloca no final se não encontrar)
            self.mensagem.emit("Inserindo tabela na página solicitada...")
            try:
                self._inserir_tabela_na_pagina(doc, tabela, self.pagina_destino)
            except Exception as e:
                self.mensagem.emit(f"Aviso: inserir na página falhou, adicionando ao final ({e})")
                doc.add_page_break()
                doc._body._body.append(tabela._tbl)

            # Salvar docx
            output_path = os.path.splitext(self.template_path)[0] + "_RELATORIO.docx"
            self.mensagem.emit("Salvando documento...")
            doc.save(output_path)

            # Converter para pdf (opcional)
            if self.converter_pdf:
                try:
                    self.mensagem.emit("Convertendo para PDF...")
                    convert(output_path)
                except Exception as e:
                    self.mensagem.emit(f"Aviso: conversão para PDF falhou: {e}")

            self.completo.emit(output_path)
        except Exception as e:
            self.completo.emit(f"ERRO: {str(e)}")

    def _inserir_tabela_na_pagina(self, doc, tabela, pagina_alvo):
        # Busca quebras de página por parágrafo (approx)
        if pagina_alvo <= 1:
            doc.paragraphs[0].insert_paragraph_before("")._p.addnext(tabela._tbl)
            return True

        paginas = 1
        for p in doc.paragraphs:
            xml = p._element.xml
            if "<w:br w:type=\"page\"" in xml or "lastRenderedPageBreak" in xml:
                paginas += 1
            if paginas == pagina_alvo:
                p._p.addnext(tabela._tbl)
                return True
        # fallback: append at end
        doc.add_page_break()
        doc.paragraphs[-1]._p.addnext(tabela._tbl)
        return True

# -----------------------------
# Interface principal
# -----------------------------
class App(QWidget):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Gerador de Relatórios LTs Perdas Duplas")
        self.resize(900, 600)

        self.template_path = None
        self.excel_path = None
        self.df = None
        self.col_widths_map = {}

        # layout principal
        main_layout = QHBoxLayout(self)

        # Preview (WebEngine)
        self.preview = QWebEngineView()
        main_layout.addWidget(self.preview, 3)

        # Painel direito
        right = QVBoxLayout()
        main_layout.addLayout(right, 1)

        # Template selector
        box_template = QGroupBox("Template Word")
        lt = QHBoxLayout()
        self.lbl_template = QLabel("Nenhum arquivo selecionado")
        btn_template = QPushButton("Selecionar")
        btn_template.clicked.connect(self.selecionar_template)
        lt.addWidget(self.lbl_template)
        lt.addWidget(btn_template)
        box_template.setLayout(lt)
        right.addWidget(box_template)

        # Excel selector
        box_excel = QGroupBox("Planilha Excel")
        lx = QHBoxLayout()
        self.lbl_excel = QLabel("Nenhum arquivo selecionado")
        btn_excel = QPushButton("Selecionar")
        btn_excel.clicked.connect(self.selecionar_excel)
        lx.addWidget(self.lbl_excel)
        lx.addWidget(btn_excel)
        box_excel.setLayout(lx)
        right.addWidget(box_excel)

        # Página destino
        right.addWidget(QLabel("Inserir tabela na página:"))
        self.spin_pagina = QSpinBox()
        self.spin_pagina.setMinimum(1)
        self.spin_pagina.setValue(1)
        right.addWidget(self.spin_pagina)

        # Style dropdown (quando template selecionado, popula)
        right.addWidget(QLabel("Estilo de Tabela (Word):"))
        self.combo_style = QComboBox()
        right.addWidget(self.combo_style)

        # Scroll area para controles de coluna
        right.addWidget(QLabel("Ajuste de Larguras por Coluna (px):"))
        scroll = QScrollArea()
        self.form_widget = QWidget()
        self.form_layout = QFormLayout()
        self.form_widget.setLayout(self.form_layout)
        scroll.setWidgetResizable(True)
        scroll.setWidget(self.form_widget)
        right.addWidget(scroll, 1)

        # Progress and controls
        self.progress = QProgressBar()
        right.addWidget(self.progress)

        btn_preview = QPushButton("Atualizar Preview")
        btn_preview.clicked.connect(self.atualizar_preview)
        right.addWidget(btn_preview)

        btn_warn = QPushButton("Aviso sobre Formatação")
        btn_warn.clicked.connect(self.mostrar_aviso)
        right.addWidget(btn_warn)

        btn_run = QPushButton("Gerar Relatório (DOCX + opcional PDF)")
        btn_run.clicked.connect(self.executar)
        right.addWidget(btn_run)

        # footer
        right.addStretch()

    def mostrar_aviso(self):
        QMessageBox.warning(
            self,
            "Atenção",
            "A formatação no Word pode não ficar exatamente como visualizado no preview.\n"
            "Ajuste larguras e estilo antes de gerar. Deseja continuar?",
            QMessageBox.Ok
        )

    def selecionar_template(self):
        path, _ = QFileDialog.getOpenFileName(self, "Template Word", "", "Documentos Word (*.docx)")
        if path:
            self.template_path = path
            self.lbl_template.setText(os.path.basename(path))
            # popula estilos
            estilos = listar_estilos_de_tabela(path)
            self.combo_style.clear()
            self.combo_style.addItems(estilos)

    def selecionar_excel(self):
        path, _ = QFileDialog.getOpenFileName(self, "Excel", "", "Planilhas Excel (*.xlsx *.xls)")
        if path:
            self.excel_path = path
            self.lbl_excel.setText(os.path.basename(path))
            self.df = pd.read_excel(self.excel_path)
            self._criar_controles_de_coluna()
            self.atualizar_preview()

    def _criar_controles_de_coluna(self):
        # limpa form_layout
        while self.form_layout.count():
            item = self.form_layout.takeAt(0)
            w = item.widget()
            if w:
                w.deleteLater()
        self.col_widths_map = {}
        if self.df is None:
            return
        for col in self.df.columns:
            spin = QSpinBox()
            spin.setRange(20, 1200)
            spin.setValue(140)
            spin.valueChanged.connect(self._on_col_width_change)
            self.form_layout.addRow(QLabel(str(col)), spin)
            self.col_widths_map[str(col)] = spin.value()

    @Slot()
    def _on_col_width_change(self):
        # atualiza map
        for i in range(self.form_layout.rowCount()):
            label_widget = self.form_layout.itemAt(i, QFormLayout.LabelRole).widget()
            field_widget = self.form_layout.itemAt(i, QFormLayout.FieldRole).widget()
            if label_widget and field_widget:
                self.col_widths_map[str(label_widget.text())] = field_widget.value()
        # opcional: atualizar preview em tempo real
        self.atualizar_preview()

    def atualizar_preview(self):
        if self.df is None:
            self.preview.setHtml("<h3>Nenhuma planilha carregada</h3>")
            return
        html = gerar_preview_html(self.df, widths=self.col_widths_map)
        self.preview.setHtml(html)

    def executar(self):
        if not self.template_path or not self.excel_path:
            QMessageBox.warning(self, "Erro", "Selecione template e planilha.")
            return

        # Modal confirmando possível diferença de formatação
        resposta = QMessageBox.question(
            self,
            "Confirmar geração",
            "A formatação pode não ficar exatamente como no preview. Deseja continuar?",
            QMessageBox.Yes | QMessageBox.No
        )
        if resposta != QMessageBox.Yes:
            return

        incluir_teste = True
        converter_pdf = True

        # prepara map de larguras atual (valores dos spinboxes)
        widths_map = {}
        for i in range(self.form_layout.rowCount()):
            label_widget = self.form_layout.itemAt(i, QFormLayout.LabelRole).widget()
            field_widget = self.form_layout.itemAt(i, QFormLayout.FieldRole).widget()
            if label_widget and field_widget:
                widths_map[str(label_widget.text())] = field_widget.value()

        table_style = self.combo_style.currentText() if self.combo_style.currentText() else "Table Grid"

        self.worker = Worker(
            self.excel_path,
            self.template_path,
            self.spin_pagina.value(),
            incluir_teste,
            converter_pdf,
            widths_map,
            table_style
        )
        self.worker.progresso.connect(self.progress.setValue)
        self.worker.mensagem.connect(lambda m: print("[INFO]", m))
        self.worker.completo.connect(self.fim)

        self.worker.start()

    def fim(self, msg):
        QMessageBox.information(self, "Finalizado", str(msg))
        self.progress.setValue(0)

# Entrypoint
if __name__ == "__main__":
    # Necessário para inicializar o WebEngine adequadamente
    app = QApplication(sys.argv)
    janela = App()
    janela.show()
    sys.exit(app.exec())
