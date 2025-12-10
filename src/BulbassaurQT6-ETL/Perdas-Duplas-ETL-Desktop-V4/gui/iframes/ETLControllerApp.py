# RelatorioWidgetFrame.py
import sys
import os
import pandas as pd
from docx import Document
from docx.shared import RGBColor, Cm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx2pdf import convert

from PySide6.QtWidgets import (
    QApplication, QFrame, QWidget, QVBoxLayout, QHBoxLayout, QLabel, QLineEdit,
    QPushButton, QTextEdit, QFileDialog, QMessageBox, QSpinBox, QGroupBox,
    QTabWidget, QTableWidget, QTableWidgetItem, QHeaderView, QFormLayout,
    QScrollArea, QComboBox
)
from PySide6.QtCore import Qt, QThread, Signal, Slot, QUrl
from PySide6.QtGui import QFont

# WebEngine (opcional)
try:
    from PySide6.QtWebEngineWidgets import QWebEngineView, QWebEngineSettings
    HAS_WEBENGINE = True
except Exception:
    HAS_WEBENGINE = False

# -----------------------
# Utilitários docx
# -----------------------
def set_cell_shading(cell, fill_hex: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def set_run_color(run, hex_color: str):
    r = int(hex_color[0:2], 16)
    g = int(hex_color[2:4], 16)
    b = int(hex_color[4:6], 16)
    run.font.color.rgb = RGBColor(r, g, b)

def px_to_cm(px: int):
    return px / 37.8

def gerar_preview_html(df, widths=None, header_fill="4472C4", header_text="FFFFFF"):
    bg = f"#{header_fill}"
    color = f"#{header_text}"
    html = '<html><head><meta charset="utf-8"></head><body>'
    html += '<table border="1" cellspacing="0" cellpadding="6" style="border-collapse:collapse;font-family:Arial;width:100%;">'
    html += '<tr>'
    for col in df.columns:
        w = widths.get(col, 140) if widths else 140
        html += f'<th style="background:{bg};color:{color};width:{w}px;">{col}</th>'
    html += '</tr>'
    for _, row in df.iterrows():
        html += '<tr>'
        for val in row:
            html += f'<td>{"" if pd.isna(val) else val}</td>'
        html += '</tr>'
    html += '</table></body></html>'
    return html

# -----------------------
# Workers
# -----------------------
class AnaliseWorker(QThread):
    finished = Signal(dict)
    error = Signal(str)

    def __init__(self, template_path, excel_path, sheet_name):
        super().__init__()
        self.template = template_path
        self.excel = excel_path
        self.sheet = sheet_name

    def run(self):
        try:
            res = {}
            if os.path.exists(self.excel):
                try:
                    df = pd.read_excel(self.excel, sheet_name=self.sheet)
                except Exception:
                    df = pd.read_excel(self.excel)
                df = df.fillna("")
                res['df'] = df
                res['rows'], res['cols'] = df.shape
            if os.path.exists(self.template):
                # contar páginas por quebras simples
                doc = Document(self.template)
                paginas = 1
                for p in doc.paragraphs:
                    xml = p._element.xml
                    if "<w:br w:type=\"page\"" in xml or "lastRenderedPageBreak" in xml:
                        paginas += 1
                        continue
                    for run in p.runs:
                        if "w:br" in run._element.xml and "type=\"page\"" in run._element.xml:
                            paginas += 1
                            break
                res['pages'] = paginas
            self.finished.emit(res)
        except Exception as e:
            self.error.emit(str(e))

class GeracaoWorker(QThread):
    finished = Signal(str)
    log = Signal(str)
    error = Signal(str)

    def __init__(self, cfg):
        super().__init__()
        self.cfg = cfg

    def run(self):
        try:
            self.log.emit("⏳ Carregando dados...")
            df = pd.read_excel(self.cfg['excel'], sheet_name=self.cfg.get('sheet', None))
            doc = Document(self.cfg['template'])

            self.log.emit("🔧 Criando tabela...")
            table = doc.add_table(rows=df.shape[0]+1, cols=df.shape[1])
            table.style = 'Table Grid'

            # cabeçalho com estilo selecionado
            sel_style = self.cfg.get('style', 'ONS')
            fill = "4472C4" if sel_style == "ONS Azul" else "A9B3BD"  # cinza
            textc = "FFFFFF"

            for j, col in enumerate(df.columns):
                cell = table.cell(0, j)
                run = cell.paragraphs[0].add_run(str(col))
                run.bold = True
                set_cell_shading(cell, fill)
                set_run_color(run, textc)

            self.log.emit("📝 Preenchendo dados...")
            for i, row in enumerate(df.values):
                for j, val in enumerate(row):
                    table.cell(i+1, j).text = "" if pd.isna(val) else str(val)

            # aplicar larguras
            self.log.emit("📐 Aplicando larguras por coluna...")
            widths_map = self.cfg.get('widths', {})
            try:
                for j, col in enumerate(df.columns):
                    px = int(widths_map.get(str(col), 140))
                    cm_val = px_to_cm(px)
                    # aplicar em cada célula da coluna
                    for r in range(df.shape[0] + 1):
                        try:
                            table.cell(r, j).width = Cm(cm_val)
                        except Exception:
                            pass
                    try:
                        table.columns[j].width = Cm(cm_val)
                    except Exception:
                        pass
            except Exception as e:
                self.log.emit(f"Aviso largura: {e}")

            # inserir na página alvo (fallback: final)
            self.log.emit("📄 Inserindo tabela no documento...")
            target_page = int(self.cfg.get('page', 1))
            try:
                if target_page <= 1:
                    doc.paragraphs[0].insert_paragraph_before("")._p.addnext(table._tbl)
                else:
                    paginas = 1
                    inserted = False
                    for p in doc.paragraphs:
                        xml = p._element.xml
                        if "<w:br w:type=\"page\"" in xml or "lastRenderedPageBreak" in xml:
                            paginas += 1
                        if paginas == target_page:
                            p._p.addnext(table._tbl)
                            inserted = True
                            break
                    if not inserted:
                        doc.add_page_break()
                        doc.paragraphs[-1]._p.addnext(table._tbl)
            except Exception as e:
                self.log.emit(f"Aviso inserção página: {e}")
                doc.add_page_break()
                doc.paragraphs[-1]._p.addnext(table._tbl)

            out_word = self.cfg.get('out_word') or os.path.splitext(self.cfg['template'])[0] + "_RELATORIO.docx"
            self.log.emit("💾 Salvando DOCX...")
            doc.save(out_word)

            if self.cfg.get('to_pdf', True):
                out_pdf = self.cfg.get('out_pdf') or os.path.splitext(out_word)[0] + ".pdf"
                try:
                    self.log.emit("🔄 Convertendo para PDF...")
                    convert(out_word, out_pdf)
                    self.finished.emit(out_pdf)
                except Exception as e:
                    self.log.emit(f"Aviso conversão: {e}")
                    self.finished.emit(out_word)
            else:
                self.finished.emit(out_word)

        except Exception as e:
            self.error.emit(str(e))

# -----------------------
# Componentes UI
# -----------------------
class ExcelPreviewTable(QTableWidget):
    def __init__(self):
        super().__init__()
        self.setAlternatingRowColors(True)
        self.setStyleSheet("""
            QTableWidget { gridline-color: #ccc; font-family: 'Segoe UI'; font-size: 13px; }
            QHeaderView::section { background-color: #005CA9; color: white; font-weight: bold; padding: 5px; border: 1px solid #00407a; }
            QTableWidget::item { padding: 5px; }
        """)

    def load_dataframe(self, df):
        self.setRowCount(df.shape[0])
        self.setColumnCount(df.shape[1])
        self.setHorizontalHeaderLabels([str(c) for c in df.columns])
        for r in range(df.shape[0]):
            for c in range(df.shape[1]):
                val = str(df.iloc[r, c])
                item = QTableWidgetItem(val)
                item.setFlags(Qt.ItemIsEnabled | Qt.ItemIsSelectable)
                self.setItem(r, c, item)
        self.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)

class PDFViewer(QWidget):
    def __init__(self):
        super().__init__()
        layout = QVBoxLayout(self)
        layout.setContentsMargins(0,0,0,0)
        if HAS_WEBENGINE:
            self.web_view = QWebEngineView()
            self.web_view.settings().setAttribute(QWebEngineSettings.PluginsEnabled, True)
            self.web_view.settings().setAttribute(QWebEngineSettings.PdfViewerEnabled, True)
            layout.addWidget(self.web_view)
        else:
            lbl = QLabel("Visualizador PDF não disponível (instale PySide6-WebEngine).")
            lbl.setAlignment(Qt.AlignCenter)
            layout.addWidget(lbl)
            self.web_view = None

    def load_pdf(self, path):
        if self.web_view and os.path.exists(path):
            self.web_view.setUrl(QUrl.fromLocalFile(os.path.abspath(path)))

# -----------------------
# FRAME (IFRAME) PRINCIPAL
# -----------------------
class RelatorioWidgetFrame(QFrame):
    """
    QFrame pronto para embed no seu UI_MainWindow.
    Importe e use: widget = RelatorioWidgetFrame(parent); main_layout.addWidget(widget)
    """
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setObjectName("RelatorioWidgetFrame")
        self.setStyleSheet("#RelatorioWidgetFrame { background: white; }")
        self._init_vars()
        self._init_ui()

    def _init_vars(self):
        self.template_path = ""
        self.excel_path = ""
        self.out_word = ""
        self.out_pdf = ""
        self.df = None
        self.col_widths = {}

    def _init_ui(self):
        main = QVBoxLayout(self)
        main.setContentsMargins(6,6,6,6)

        tabs = QTabWidget()
        tabs.setStyleSheet("""
            QTabBar::tab { padding: 8px 14px; }
            QTabBar::tab:selected { background: #005CA9; color: white; }
        """)
        self.tabs = tabs

        # Aba Controle (com seu layout)
        self.tab_control = QWidget()
        self._build_control_tab()
        # Aba de preview Excel
        self.tab_excel = ExcelPreviewTable()
        # Aba PDF viewer
        self.tab_pdf = PDFViewer()

        tabs.addTab(self.tab_control, "🎛️ Painel de Controle")
        tabs.addTab(self.tab_excel, "📊 Prévia dos Dados")
        tabs.addTab(self.tab_pdf, "📄 Visualizador PDF")

        main.addWidget(tabs)

    def _build_control_tab(self):
        container = QWidget()
        h = QHBoxLayout(container)
        h.setContentsMargins(0,0,0,0)

        left = QVBoxLayout()
        left.setSpacing(8)

        left.addWidget(QLabel("📂 Entradas"))
        self.le_template = self._file_input("Template Word (.docx):", left, save=False)
        self.le_excel = self._file_input("Planilha Excel (.xlsx):", left, save=False)
        sheet_h = QHBoxLayout()
        sheet_h.addWidget(QLabel("Aba:"))
        self.le_sheet = QLineEdit("Sheet1")
        sheet_h.addWidget(self.le_sheet)
        left.addLayout(sheet_h)

        left.addSpacing(6)
        left.addWidget(QLabel("💾 Saídas"))
        self.le_out_word = self._file_input("Salvar Word (.docx):", left, save=True)
        self.le_out_pdf = self._file_input("Salvar PDF (.pdf):", left, save=True)

        left.addSpacing(10)

        # pagina spin
        pg_box = QGroupBox("Onde inserir a tabela")
        pg_layout = QHBoxLayout(pg_box)
        pg_layout.addWidget(QLabel("Página:"))
        self.spin_page = QSpinBox()
        self.spin_page.setRange(1, 999)
        pg_layout.addWidget(self.spin_page)
        left.addWidget(pg_box)

        left.addWidget(QLabel("Estilo de Cabeçalho:"))
        self.combo_style = QComboBox()
        self.combo_style.addItems(["ONS Azul", "Cinza Corporativo"])
        left.addWidget(self.combo_style)

        left.addWidget(QLabel("Ajuste de larguras (px)"))
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        self.form = QWidget()
        self.form_layout = QFormLayout()
        self.form.setLayout(self.form_layout)
        scroll.setWidget(self.form)
        left.addWidget(scroll, 1)

        btn_analyze = QPushButton("🔍 Analisar")
        btn_analyze.clicked.connect(self.analisar)
        left.addWidget(btn_analyze)

        btn_preview = QPushButton("🔄 Atualizar Preview")
        btn_preview.clicked.connect(self.atualizar_preview)
        left.addWidget(btn_preview)

        btn_warn = QPushButton("⚠️ Aviso formatação")
        btn_warn.clicked.connect(self._show_warning)
        left.addWidget(btn_warn)

        btn_generate = QPushButton("🚀 Gerar Relatório")
        btn_generate.clicked.connect(self.gerar)
        left.addWidget(btn_generate)

        left.addStretch()

        right = QVBoxLayout()
        right.addWidget(QLabel("📟 Console"))
        self.console = QTextEdit()
        self.console.setReadOnly(True)
        self.console.setStyleSheet("background:#0f0f0f;color:#00ff81;font-family: Consolas;")
        right.addWidget(self.console)

        h.addLayout(left, 0)
        h.addLayout(right, 1)

        self.tab_control.setLayout(h)

    def _file_input(self, label_text, layout, save=False):
        layout.addWidget(QLabel(label_text))
        h = QHBoxLayout()
        le = QLineEdit()
        btn = QPushButton("...")
        if save:
            btn.clicked.connect(lambda: self._browse_save(le))
        else:
            btn.clicked.connect(lambda: self._browse_open(le))
        h.addWidget(le)
        h.addWidget(btn)
        layout.addLayout(h)
        return le

    def _browse_open(self, le):
        f, _ = QFileDialog.getOpenFileName(self, "Abrir")
        if f: le.setText(f)

    def _browse_save(self, le):
        f, _ = QFileDialog.getSaveFileName(self, "Salvar")
        if f: le.setText(f)

    def _show_warning(self):
        QMessageBox.warning(self, "Atenção", "A formatação pode não ficar exatamente como no preview. Ajuste larguras e estilo antes de gerar.")

    # --- Ações ---
    def analisar(self):
        tpl = self.le_template.text()
        xls = self.le_excel.text()
        sheet = self.le_sheet.text() or None
        if not tpl or not xls:
            QMessageBox.warning(self, "Erro", "Informe template e planilha.")
            return
        self.console.append(">> Iniciando análise...")
        self.worker_analise = AnaliseWorker(tpl, xls, sheet)
        self.worker_analise.finished.connect(self._pos_analise)
        self.worker_analise.error.connect(lambda e: QMessageBox.critical(self, "Erro análise", e))
        self.worker_analise.start()

    def _pos_analise(self, res):
        self.console.append(">> Análise finalizada.")
        if 'df' in res:
            self.df = res['df']
            self.tab_excel.load_dataframe(self.df)
            self._build_width_controls()
            self.console.append(f">> Dados: {res.get('rows',0)} linhas, {res.get('cols',0)} colunas.")
        if 'pages' in res:
            self.spin_page.setMaximum(res['pages'] + 1)
            self.console.append(f">> Documento com ~{res['pages']} páginas.")

    def _build_width_controls(self):
        # limpa
        while self.form_layout.count():
            item = self.form_layout.takeAt(0)
            w = item.widget()
            if w:
                w.deleteLater()
        self.col_widths = {}
        if self.df is None:
            return
        for col in self.df.columns:
            spin = QSpinBox()
            spin.setRange(20, 1200)
            spin.setValue(140)
            spin.valueChanged.connect(self.atualizar_preview)
            self.form_layout.addRow(QLabel(str(col)), spin)
            self.col_widths[str(col)] = spin.value()

    @Slot()
    def atualizar_preview(self):
        # atualiza mapa de larguras
        if self.df is None:
            return
        for i in range(self.form_layout.rowCount()):
            label_widget = self.form_layout.itemAt(i, QFormLayout.LabelRole).widget()
            field_widget = self.form_layout.itemAt(i, QFormLayout.FieldRole).widget()
            if label_widget and field_widget:
                self.col_widths[str(label_widget.text())] = field_widget.value()
        # gerar html e mostrar na aba Excel (se WebEngine disponível, abrir HTML simples no PDF tab)
        html = gerar_preview_html(self.df, widths=self.col_widths,
                                 header_fill="4472C4" if self.combo_style.currentText()=="ONS Azul" else "A9B3BD")
        if HAS_WEBENGINE:
            # mostra o preview HTML convertendo para um file:// temporário
            tmp = os.path.join(os.path.abspath(os.path.dirname(__file__)), "preview_temp.html")
            with open(tmp, "w", encoding="utf-8") as f:
                f.write(html)
            # usa webengine no tab_pdf se disponível
            if self.tab_pdf.web_view:
                self.tab_pdf.web_view.setUrl(QUrl.fromLocalFile(tmp))
                self.tabs.setCurrentWidget(self.tab_pdf)
        else:
            # carrega tabela no QTableWidget (aba excel)
            self.tab_excel.load_dataframe(self.df)

    def gerar(self):
        tpl = self.le_template.text()
        xls = self.le_excel.text()
        if not tpl or not xls:
            QMessageBox.warning(self, "Erro", "Informe template e planilha.")
            return
        # confirmar aviso
        resp = QMessageBox.question(self, "Confirmar", "A formatação pode não ser idêntica ao preview. Deseja continuar?")
        if resp != QMessageBox.StandardButton.Yes:
            return
        cfg = {
            'template': tpl,
            'excel': xls,
            'sheet': self.le_sheet.text() or None,
            'out_word': self.le_out_word.text() or os.path.splitext(tpl)[0] + "_RELATORIO.docx",
            'out_pdf': self.le_out_pdf.text() or os.path.splitext(tpl)[0] + "_RELATORIO.pdf",
            'page': self.spin_page.value(),
            'style': self.combo_style.currentText(),
            'widths': {k: (self.form_layout.itemAt(i, QFormLayout.FieldRole).widget().value()
                         if self.form_layout.itemAt(i, QFormLayout.FieldRole).widget() else 140)
                       for i, k in enumerate([str(c) for c in (self.df.columns if self.df is not None else [])])},
            'to_pdf': True
        }
        self.console.append(">> Iniciando geração...")
        self.worker_gen = GeracaoWorker(cfg)
        self.worker_gen.log.connect(self.console.append)
        self.worker_gen.finished.connect(self._pos_geracao)
        self.worker_gen.error.connect(lambda e: QMessageBox.critical(self, "Erro geração", e))
        self.worker_gen.start()

    def _pos_geracao(self, out_path):
        self.console.append(f">> Geração finalizada: {out_path}")
        # tenta carregar pdf no viewer
        if os.path.exists(out_path) and out_path.lower().endswith(".pdf"):
            self.tab_pdf.load_pdf(out_path)
            self.tabs.setCurrentWidget(self.tab_pdf)
            QMessageBox.information(self, "Sucesso", "Relatório gerado e aberto no visualizador.")
        else:
            QMessageBox.information(self, "Finalizado", f"Arquivo gerado: {out_path}")

# Testa de forma isolada
if __name__ == "__main__":
    app = QApplication(sys.argv)
    app.setFont(QFont("Segoe UI", 10))
    frame = RelatorioWidgetFrame()
    frame.setWindowTitle("Iframe - Relatório (teste isolado)")
    frame.resize(1100, 700)
    frame.show()
    sys.exit(app.exec())
