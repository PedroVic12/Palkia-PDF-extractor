import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.shared import Inches
import os
from docx2pdf import convert
from tqdm import tqdm
import time

def adicionar_pagina_e_tabela(doc_path: str, excel_path: str, output_docx_path: str):
    print(f"Abrindo documento Word template: {doc_path}")
    doc = Document(doc_path)
    
    # Adiciona a nova seção e configura margens
    nova_secao = doc.add_section(WD_SECTION.NEW_PAGE)
    nova_secao.top_margin = Inches(1)
    nova_secao.bottom_margin = Inches(1)
    nova_secao.left_margin = Inches(1)
    nova_secao.right_margin = Inches(1)

    doc.add_heading('Relatório de Perdas Duplas (Dados do Excel)', level=1)
    doc.add_paragraph("Abaixo estão os dados extraídos da planilha Excel:")

    # Carrega os dados do Excel
    df = pd.read_excel(excel_path)
    rows, cols = df.shape
    table = doc.add_table(rows + 1, cols)
    table.style = 'Table Grid'

    # Preenche o cabeçalho da tabela
    for j, column_name in enumerate(df.columns):
        table.cell(0, j).text = str(column_name)

    # Preenche os dados com barra de progresso (tqdm)
    print(f"Preenchendo tabela com {rows} linhas...")
    for i in tqdm(range(rows), desc="Progresso da Tabela"):
        for j in range(cols):
            table.cell(i + 1, j).text = str(df.values[i, j])

    doc.save(output_docx_path)
    print(f"Documento Word modificado salvo em '{output_docx_path}'.")
    return output_docx_path

def converter_docx_para_pdf(docx_path: str, pdf_path: str):
    print(f"Iniciando conversão de DOCX para PDF. Isso pode demorar MUITO...")
    start_time = time.time()
    
    try:
        # A conversão é uma operação bloqueadora. Não há progresso iterativo aqui.
        convert(docx_path, pdf_path)
        end_time = time.time()
        print(f"Conversão concluída em {end_time - start_time:.2f} segundos.")
        print(f"Arquivo PDF salvo em '{pdf_path}'.")
    except Exception as e:
        print(f"ERRO na conversão para PDF. Certifique-se que o Microsoft Office está instalado.")
        print(f"Detalhes do erro: {e}")
        raise

# --- Execução Principal ---
if __name__ == "__main__":
    template_word_path = r"C:\Users\pedrovictor.veras\OneDrive - Operador Nacional do Sistema Eletrico\Documentos\ESTAGIO_ONS_PVRV_2025\GitHub\Palkia-PDF-extractor\src\BulbassaurQT6-ETL\sistema-ferramentas-RPA-desktop\app\assets\docs\Lista de Contingências Duplas - Copia.docx" 
    excel_data_path = r"C:\Users\pedrovictor.veras\OneDrive - Operador Nacional do Sistema Eletrico\Documentos\ESTAGIO_ONS_PVRV_2025\GitHub\Palkia-PDF-extractor\src\BulbassaurQT6-ETL\sistema-ferramentas-RPA-desktop\app\assets\planilhas_PLC\perdas_duplas_ETL_corrigido.xlsx"
    
    output_word_path = "documento_final_gerado.docx"
    output_pdf_path = "relatorio_PLC_perdas_duplas.pdf"

    if os.path.exists(template_word_path) and os.path.exists(excel_data_path):
        try:
            # Etapa 1: Gerar DOCX (com barra de progresso na tabela)
            docx_gerado_path = adicionar_pagina_e_tabela(
                template_word_path, 
                excel_data_path, 
                output_word_path
            )
            
            # Etapa 2: Converter para PDF (Monitoramento de tempo)
            converter_docx_para_pdf(docx_gerado_path, output_pdf_path)
            
        except Exception as e:
            print(f"\nO processamento principal falhou.")

    else:
        print("\nERRO CRÍTICO: Um ou ambos os arquivos de entrada não foram encontrados.")

